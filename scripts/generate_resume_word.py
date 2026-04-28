"""One-off generator: Senior Technology Architect resume -> .docx"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUTPUT = Path(r"d:\Users\Krishan.Tiwari\Downloads\Krishan-Kumar-Tiwari-Senior-Technology-Architect-Resume.docx")

SKILLS = [
    "Microservices Architecture",
    "Event-Driven Architecture",
    "RESTful & GraphQL APIs",
    "Java / Spring Boot",
    "React.js",
    "AWS",
    "Microsoft Azure",
    "IBM Sterling OMS",
    "SAP B1 / D365",
    "SQL Server / MongoDB",
    "Azure IoT Hub",
    "Dynamics Navision ERP",
    "Fluent OMS",
    "Power BI",
    "Snowflake",
    "C++ / SQL",
]

EXPERIENCE = [
    {
        "title": "Technology Architect",
        "company": "MAIR Group",
        "period": "Jan 2025 – Present",
        "location": "Abu Dhabi, UAE",
        "bullets": [
            "Architected middleware integrating Microsoft D365 with Electronic Shelf Label (ESL) hardware, enabling real-time automated price synchronisation across all retail touchpoints.",
            "Designed and delivered D365 + SAP ERP middleware with Azure validation layer for price, item, barcode, and sales data synchronisation across enterprise systems.",
            "Led end-to-end solution design, stakeholder alignment, and delivery governance for multi-vendor enterprise integrations.",
        ],
    },
    {
        "title": "Manager – Technology",
        "company": "GMG International Pvt. Ltd.",
        "period": "Apr 2024 – Dec 2024",
        "location": "India",
        "bullets": [
            "Built and managed middleware connecting daily mall sales feeds into GMG core systems, automating reporting and accelerating onboarding of new mall integrations.",
            "Delivered Amazon Rush integration between GMG SAP and Amazon's fulfilment platform — automating inventory, price-list exchange, and sales-order processing with direct D365 invoice generation.",
            "Managed cross-functional delivery teams and defined architectural standards for middleware services.",
        ],
    },
    {
        "title": "Project Lead",
        "company": "Mphasis Ltd.",
        "period": "Sep 2023 – Apr 2024",
        "location": "India",
        "bullets": [
            "Led enterprise architecture and delivery of large-scale integration projects for global clients.",
            "Defined technical roadmap and led architecture reviews, ensuring alignment with business objectives.",
            "Mentored senior engineers, established coding standards, and drove adoption of best practices.",
        ],
    },
    {
        "title": "Team Lead",
        "company": "TechVillage Soft Pvt. Ltd.",
        "period": "Oct 2022 – Sep 2023",
        "location": "India",
        "bullets": [
            "Designed Fleet Management System for Serh Group — covering CRM (Lead Management, Proposals), Project Management (MSA, Work Orders), HR, Workshop, and Inventory modules with end-to-end operational tracking from customer onboarding to asset lifecycle.",
            "Integrated Fleet Management System with SAP B1 and an Asset Tagging platform for real-time data synchronisation and reporting.",
            "Architected middleware connecting AFMS with SAP and Snowflake; executed asset tagging and testing at live oil & gas sites.",
        ],
    },
    {
        "title": "Senior Software Engineer",
        "company": "Marks & Spencer Reliance India Pvt. Ltd.",
        "period": "Feb 2021 – Oct 2022",
        "location": "Gurugram, India",
        "bullets": [
            "Led migration of Fluent OMS to IBM Sterling OMS on AWS (MNS OMS 2.0) — developed services, managed configurations, and extended tables for a high-volume retail order-management platform.",
            "Integrated Fluent OMS with IBM Watson IVR; developed GraphQL queries and Java services for order processing and fulfillment reporting (MNS OMS 1.0).",
        ],
    },
    {
        "title": "Software Engineer",
        "company": "Motherson Sumi Infotech & Design (MSID)",
        "period": "May 2017 – Mar 2020",
        "location": "Noida, India",
        "bullets": [
            "Built the DAART Audit Platform — a configurable Java Rule Engine for vendor data standardisation and auditing, paired with Power BI dashboards monitoring KPIs and exceptions.",
            "Developed live IoT OEE dashboards using Azure IoT Hub and Stream Analytics, integrating multiple ERP data sources into retrospective production-performance dashboards.",
        ],
    },
]

PROJECTS = [
    {
        "name": "Fleet Management System",
        "client": "Serh Group",
        "stack": ["Java", "Spring Boot", "SAP B1", "Snowflake", "REST APIs"],
        "summary": "End-to-end operational platform covering CRM, project management, fleet, HR, workshop, and inventory with real-time SAP B1 and asset tagging integration.",
    },
    {
        "name": "ERP Middleware (D365 + SAP)",
        "client": "Mair Group",
        "stack": ["Azure", "D365", "SAP", "REST APIs", "Java"],
        "summary": "Middleware layer synchronising price, item, barcode, and sales data between Microsoft D365 and SAP with Azure-based validation, enabling near-real-time ERP coherence.",
    },
    {
        "name": "Electronic Shelf Label (ESL)",
        "client": "Mair Group",
        "stack": ["D365", "ESL Hardware", "Spring Boot", "REST APIs"],
        "summary": "Automated price-update pipeline from D365 to retail ESL hardware, eliminating manual intervention and ensuring 100% accuracy at shelf.",
    },
    {
        "name": "MNS OMS 2.0 — IBM Sterling Migration",
        "client": "Marks & Spencer Reliance",
        "stack": ["IBM Sterling OMS", "AWS", "Java", "Spring Boot"],
        "summary": "Full migration from Fluent OMS to IBM Sterling OMS hosted on AWS for a high-volume omnichannel retail order platform serving millions of customers.",
    },
    {
        "name": "Amazon Rush Integration",
        "client": "GMG International",
        "stack": ["SAP", "D365", "Amazon APIs", "Java"],
        "summary": "Automated inventory, price-list, and sales-order exchange between GMG SAP and Amazon Rush with direct invoice generation into D365.",
    },
    {
        "name": "IoT OEE Dashboard",
        "client": "Motherson Sumi (MSID)",
        "stack": ["Azure IoT Hub", "Stream Analytics", "Power BI", "ERP APIs"],
        "summary": "Real-time manufacturing OEE dashboards powered by Azure IoT Hub and Stream Analytics, integrated with multiple ERP sources for live and retrospective analytics.",
    },
]

STRENGTHS_ROWS = [
    ("Enterprise Integration", "ERP (SAP B1, D365, Dynamics Navision), OMS (IBM Sterling, Fluent), WMS (Uniware), IVR (Watson)"),
    ("Cloud & Infrastructure", "AWS (hosting, services), Azure (IoT Hub, Stream Analytics, validation pipelines), cloud-native design"),
    ("Backend Engineering", "Java, Spring Boot, Microservices, REST APIs, GraphQL, JDBC, event-driven patterns"),
    ("Frontend & Reporting", "React.js, Power BI dashboards, SOLR search, Hybris/WCMS content management"),
    ("Data & Databases", "SQL Server, MongoDB, Oracle SQL Developer, Snowflake, data modelling & migration"),
    ("Architecture Patterns", "Microservices, Domain-Driven Design, Event-Driven Architecture, API Gateway, CQRS"),
    ("Project Delivery", "Architecture governance, stakeholder management, team mentoring, release management"),
]

PROFILE = (
    "Accomplished Senior Technology Architect with over 10 years of hands-on experience "
    "designing, delivering, and governing enterprise-scale systems across retail, logistics, "
    "manufacturing, and oil & gas. Deep expertise in Java, Spring Boot, Microservices, and "
    "Cloud-native solutions on AWS and Azure. Proven track record architecting complex "
    "ERP, OMS, and WMS integrations — including IBM Sterling, SAP B1/D365, and Fluent OMS — "
    "for global brands such as Marks & Spencer and GMG International. Adept at bridging "
    "business strategy and technical execution, leading cross-functional teams, and "
    "establishing architectural standards that scale."
)


def set_run_font(run, size_pt: float | None = None, bold: bool = False, color: RGBColor | None = None) -> None:
    font = run.font
    if size_pt is not None:
        font.size = Pt(size_pt)
    font.bold = bold
    if color is not None:
        font.color.rgb = color


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 22, bold=True)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 12, bold=True)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)


def add_contact_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_h2(doc: Document, text: str) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 13, bold=True)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 11)


def add_stats_row(doc: Document) -> None:
    p = doc.add_paragraph()
    parts = [
        ("10+", "Years Experience"),
        ("12+", "Enterprise Projects"),
        ("AWS + Azure", "Cloud Platforms"),
        ("4+", "Countries / Clients"),
    ]
    for i, (val, lbl) in enumerate(parts):
        if i:
            p.add_run("   |   ")
        r1 = p.add_run(f"{val} ")
        set_run_font(r1, 11, bold=True)
        r2 = p.add_run(lbl)
        set_run_font(r2, 10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(54)
    sec.bottom_margin = Pt(54)
    sec.left_margin = Pt(54)
    sec.right_margin = Pt(54)

    add_h1(doc, "Krishan Kumar Tiwari")
    add_subtitle(doc, "Senior Technology Architect")
    add_contact_line(doc, "krishankmr478@gmail.com  |  +91 9821532394  |  Abu Dhabi, UAE")
    add_contact_line(doc, "linkedin.com/in/krishan-kumar-91071076")

    add_stats_row(doc)

    add_h2(doc, "Professional Profile")
    add_body(doc, PROFILE)

    add_h2(doc, "Core Competencies")
    add_body(doc, " · ".join(SKILLS))

    add_h2(doc, "Career History")
    for job in EXPERIENCE:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{job['title']} — {job['company']}")
        set_run_font(r1, 11, bold=True)
        r2 = p.add_run(f"    ({job['period']})")
        set_run_font(r2, 10)
        pl = doc.add_paragraph(job["location"])
        for r in pl.runs:
            set_run_font(r, 10)
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        for b in job["bullets"]:
            bp = doc.add_paragraph(style="List Bullet")
            run = bp.add_run(b)
            set_run_font(run, 10.5)

    add_h2(doc, "Key Projects")
    for pr in PROJECTS:
        hp = doc.add_paragraph()
        t = hp.add_run(pr["name"])
        set_run_font(t, 11, bold=True)
        hp.add_run("  —  ")
        c = hp.add_run(f"Client: {pr['client']}")
        set_run_font(c, 10)
        c.font.italic = True
        add_body(doc, pr["summary"])
        sp = doc.add_paragraph()
        sr = sp.add_run("Technologies: " + ", ".join(pr["stack"]))
        set_run_font(sr, 10)
        sp.paragraph_format.space_after = Pt(8)

    add_h2(doc, "Architecture & Technical Strengths")
    table = doc.add_table(rows=1 + len(STRENGTHS_ROWS), cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Domain"
    hdr[1].text = "Capabilities"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                set_run_font(r, 10, bold=True)
    for i, (dom, cap) in enumerate(STRENGTHS_ROWS, start=1):
        row = table.rows[i].cells
        row[0].text = dom
        row[1].text = cap
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run_font(r, 10)

    add_h2(doc, "Education")
    add_body(doc, "Master of Computer Applications (MCA) — Maharshi Dayanand University, 2012–2015. Advanced Institute of Technology & Management, Faridabad.")
    add_body(doc, "Bachelor of Computer Applications (BCA) — Maharshi Dayanand University, 2009–2012. G.G.D.S.D. College, Palwal.")

    add_h2(doc, "Target Roles")
    roles = [
        "Principal Architect",
        "Enterprise Architect",
        "Senior Solutions Architect",
        "Cloud Architect",
        "Platform Architect",
        "Distinguished Engineer",
        "Technical Program Lead",
    ]
    add_body(doc, " · ".join(roles))
    add_body(
        doc,
        "Open to opportunities at Microsoft, Amazon, Google, Salesforce, SAP, Oracle, IBM, Accenture, Deloitte, and other global technology leaders.",
    )

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("Resume prepared April 2026 · krishankmr478@gmail.com · +91 9821532394")
    set_run_font(fr, 9)
    fr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote: {OUTPUT}")


if __name__ == "__main__":
    main()
