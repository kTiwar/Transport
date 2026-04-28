# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
OUT = r"d:\Transport backup\Transport\Krishan_Kumar_Tiwari_Resume_India_ATS.docx"
OUT_ONEDRIVE = r"d:\Users\Krishan.Tiwari\OneDrive - Abu Dhabi Co-Operative Society\Desktop\file\Krishan_Kumar_Tiwari_Resume_India_ATS.docx"

def set_doc_defaults(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

def heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True

def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it.strip(), style="List Bullet")

def main():
    doc = Document()
    set_doc_defaults(doc)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("KRISHAN KUMAR TIWARI")
    r.bold = True
    r.font.size = Pt(16)
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.add_run("Senior Technology Architect | Solution and Integration Architecture").font.size = Pt(11)
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run("krishankmr478@gmail.com | +91-9821532394 | Gurugram, Haryana, India (open to PAN-India and hybrid roles)")
    ln = doc.add_paragraph()
    ln.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ln.add_run("linkedin.com/in/krishan-kumar-91071076")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("11+ years | Enterprise integration, microservices, cloud (AWS, Azure)")
    heading(doc, "Professional Summary")
    doc.add_paragraph("Senior Technology Architect with 11+ years designing and delivering enterprise systems across retail, fleet, asset management, and manufacturing. Strong in Java, Spring Boot, microservices, and cloud-native patterns on AWS and Azure. Proven delivery of ERP, OMS, and WMS integrations (IBM Sterling, SAP B1, Microsoft D365, Fluent OMS) for global and Indian operations including Marks and Spencer Reliance and GMG. Experienced leading architecture, stakeholder alignment, mentoring, and delivery governance.")
    heading(doc, "Core Skills")
    bullets(doc, ["Architecture: Microservices, event-driven design, REST, GraphQL, API gateways, DDD, CQRS", "Platforms: Java, Spring Boot, React.js, IBM Sterling OMS, Fluent OMS, SAP B1, D365, Dynamics NAV", "Cloud: AWS, Microsoft Azure, Azure IoT Hub, Stream Analytics, integration middleware", "Data: SQL Server, MongoDB, Snowflake, reporting with Power BI", "Delivery: Product ownership, agile delivery, code standards, architecture reviews"])
    heading(doc, "Professional Experience")
    def job(title, company, loc, start_end, points):
        p = doc.add_paragraph()
        p.add_run(title + " - " + company).bold = True
        p2 = doc.add_paragraph()
        p2.add_run(start_end + " | " + loc).italic = True
        bullets(doc, points)
    job("Technology Architect", "MAIR Group", "Abu Dhabi, UAE", "January 2025 - Present", ["Product owner for fleet management (CRM, projects, fleet, HR, workshop, inventory) for Serh Group; SAP B1 and asset-tagging integration.", "Product owner for asset tagging (masters, transactions, reporting) and transport management (orders, partners, routes, third-party integrations).", "Architected D365 to Electronic Shelf Label (ESL) middleware for automated in-store pricing.", "Architected Talabat and InstaShop integrations with POS and D365.", "Designed D365 and SAP ERP middleware with Azure validation for price, item, barcode, and sales-data synchronization.", "Led solution design, stakeholder alignment, and multi-vendor integration governance."])
    job("Manager, Technology", "GMG International Pvt. Ltd.", "India", "April 2024 - December 2024", ["Built middleware for daily mall sales feeds into core systems; faster onboarding of new malls.", "Delivered Amazon Rush integration (GMG SAP and Amazon fulfilment): inventory, price lists, sales orders, D365 invoicing.", "Managed cross-functional teams and middleware architecture standards."])
    job("Project Lead", "Mphasis Ltd.", "India", "September 2023 - April 2024", ["Enterprise architecture and delivery for JPMorgan Chase, including First Republic Bank integration.", "Technical roadmap, architecture reviews, mentoring, and engineering best practices."])
    job("Team Lead", "TechVillage Soft Pvt. Ltd.", "India", "October 2022 - September 2023", ["Dynamics NAV ERP integrations with marketplaces; WMS integration and support.", "Vendor portal extensions; 3PL and logistics integrations."])
    job("Senior Software Engineer", "Team Computers / Marks and Spencer Reliance India Pvt. Ltd.", "Gurugram, India", "March 2020 - October 2022", ["Led Fluent OMS to IBM Sterling OMS (MNS OMS 2.0) migration on AWS: services, configuration, schema extensions for high-volume retail.", "Fluent OMS and IBM Watson: GraphQL and Java services for order processing and fulfilment reporting (MNS OMS 1.0)."])
    job("Software Engineer", "Motherson Sumi Infotech and Design (MSID)", "Noida, India", "May 2017 - March 2020", ["Built Digital Audit and Analytics Tool (DAART) for internal audit; 30+ group companies, higher audit throughput and 30%+ efficiency gains.", "IoT OEE dashboards using Azure IoT Hub and Stream Analytics with multi-ERP integration."])
    job("Software Engineer", "Pragiti Internet Technology Pvt. Ltd.", "Noida, India", "August 2016 - May 2017", ["SAP Hybris development for Maui Jim B2C e-commerce."])
    job("Software Engineer", "Mob Fountain Media Pvt. Ltd.", "Noida, India", "February 2015 - August 2016", ["CMS development for FreeTalkTime Android application."])
    heading(doc, "Key Projects")
    bullets(doc, ["Fleet Management System (Serh): CRM, PM, fleet, HR, workshop, inventory; SAP B1 sync - Java, Spring Boot, SAP B1, Snowflake, REST.", "ERP middleware (MAIR): D365 and SAP sync with Azure validation - Azure, D365, SAP, Java, REST.", "ESL integration (MAIR): D365 to shelf hardware pricing automation - D365, Spring Boot, REST.", "MNS OMS 2.0 (M and S Reliance): Fluent to IBM Sterling on AWS - Sterling, AWS, Java, Spring Boot.", "Amazon Rush (GMG): SAP, Amazon APIs, D365 invoicing - SAP, D365, Java.", "IoT OEE (Motherson): Azure IoT Hub, Stream Analytics, Power BI, ERP APIs."])
    heading(doc, "Education")
    doc.add_paragraph("Master of Computer Applications (MCA), 2012-2015 - Maharshi Dayanand University; Advanced Institute of Technology and Management, Faridabad")
    doc.add_paragraph("Bachelor of Computer Applications (BCA), 2009-2012 - Maharshi Dayanand University; G.G.D.S.D. College, Palwal")
    doc.save(OUT)
    print("Wrote:", OUT)
    try:
        import shutil
        shutil.copy2(OUT, OUT_ONEDRIVE)
        print("Copied to OneDrive:", OUT_ONEDRIVE)
    except OSError as e:
        print("OneDrive copy skipped:", e)

if __name__ == "__main__":
    main()